"""
Document Parser — Extract text from PDF, DOCX, HTML, and plain text files.
محلل المستندات — استخراج النص من ملفات PDF وDOCX وHTML والنصوص العادية

Lazy imports: only loads libraries when needed.
Falls back gracefully if a library is missing.
"""
from pathlib import Path
from typing import Optional
from loguru import logger


def extract_text(file_path: str) -> Optional[str]:
    """
    Extract text from any supported file type.
    استخراج النص من أي نوع ملف مدعوم

    Supported: .txt, .md, .json, .pdf, .docx, .html, .csv
    Returns None if extraction fails.
    """
    path = Path(file_path)
    if not path.exists():
        logger.error(f"File not found: {file_path}")
        return None

    suffix = path.suffix.lower()
    extractors = {
        ".txt": _extract_plain,
        ".md": _extract_plain,
        ".json": _extract_plain,
        ".csv": _extract_plain,
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,
        ".html": _extract_html,
        ".htm": _extract_html,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        logger.warning(f"Unsupported file type: {suffix} ({file_path})")
        return None

    try:
        text = extractor(str(path))
        if text:
            logger.debug(f"Extracted {len(text)} chars from {path.name}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from {path.name}: {e}")
        return None


def _extract_plain(path: str) -> str:
    """Read plain text file."""
    return Path(path).read_text(encoding="utf-8")


def _extract_pdf(path: str) -> Optional[str]:
    """
    Extract text from PDF using fitz (PyMuPDF).
    Falls back to pdfplumber if fitz unavailable.
    """
    # Try PyMuPDF first (faster, better quality)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[صفحة {page_num + 1}]\n{text}")
        doc.close()
        return "\n\n".join(pages) if pages else None
    except ImportError:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[صفحة {i + 1}]\n{text}")
        return "\n\n".join(pages) if pages else None
    except ImportError:
        logger.error(
            "PDF extraction requires PyMuPDF or pdfplumber. "
            "Install: pip install PyMuPDF  or  pip install pdfplumber"
        )
        return None


def _extract_docx(path: str) -> Optional[str]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs) if paragraphs else None
    except ImportError:
        logger.error("DOCX extraction requires python-docx. Install: pip install python-docx")
        return None


def _extract_html(path: str) -> Optional[str]:
    """Extract text from HTML — strips tags."""
    import re
    raw = Path(path).read_text(encoding="utf-8")
    # Remove script/style blocks
    raw = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', raw, flags=re.DOTALL | re.IGNORECASE)
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', ' ', raw)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text if text else None


def supported_extensions() -> list[str]:
    """Return list of supported file extensions."""
    return [".txt", ".md", ".json", ".csv", ".pdf", ".docx", ".doc", ".html", ".htm"]
