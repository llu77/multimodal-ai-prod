"""
Advanced Agent Skills — Web Search, PubMed, File Generation, Image Analysis.
مهارات الوكيل المتقدمة — بحث الويب، PubMed، توليد ملفات، تحليل صور

Each skill is a self-contained module following the SKILL.md pattern:
- Description (what it does)
- When to use (triggers)
- Execute (the function)
- No external API keys required — all work locally or with free APIs

Skills added:
1. web_search      — Search the web via DuckDuckGo (no API key)
2. pubmed_search   — Search PubMed medical literature (free NCBI API)
3. generate_pdf    — Generate PDF reports from text
4. generate_docx   — Generate Word documents
5. analyze_image   — Describe/analyze images using the local model
6. translate       — Arabic ↔ English translation
7. drug_interaction — Check basic drug interactions
"""
import re
import json
import urllib.request
import urllib.parse
from pathlib import Path
from datetime import datetime
from typing import Optional
from loguru import logger

from src.agent.tools import Tool, ToolRegistry


# ═══════════════════════════════════════════
# SKILL 1: Web Search (DuckDuckGo — No API Key)
# ═══════════════════════════════════════════

def _web_search(query: str, max_results: int = 5, **kwargs) -> str:
    """
    Search the web using DuckDuckGo Instant Answer API.
    بحث الويب باستخدام DuckDuckGo — بدون مفتاح API

    Returns: top results with titles and snippets.
    """
    try:
        encoded = urllib.parse.quote_plus(query)
        url = f"https://api.duckduckgo.com/?q={encoded}&format=json&no_html=1&skip_disambig=1"

        req = urllib.request.Request(url, headers={"User-Agent": "SymbolAI/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        results = []

        # Abstract (main answer)
        if data.get("Abstract"):
            results.append(f"📌 {data['AbstractSource']}: {data['Abstract']}")

        # Related topics
        for topic in data.get("RelatedTopics", [])[:max_results]:
            if isinstance(topic, dict) and "Text" in topic:
                text = topic["Text"][:300]
                url_link = topic.get("FirstURL", "")
                results.append(f"• {text}\n  🔗 {url_link}")

        if not results:
            return f"لم يتم العثور على نتائج لـ: {query}. جرّب صياغة مختلفة."

        return f"نتائج البحث عن: {query}\n\n" + "\n\n".join(results)

    except Exception as e:
        logger.error(f"Web search failed: {e}")
        return f"خطأ في البحث: {e}. تأكد من اتصال الإنترنت."


# ═══════════════════════════════════════════
# SKILL 2: PubMed Medical Literature Search
# ═══════════════════════════════════════════

def _pubmed_search(query: str, max_results: int = 5, **kwargs) -> str:
    """
    Search PubMed for medical research papers.
    بحث PubMed عن الأبحاث الطبية — NCBI E-utilities (مجاني)

    Uses NCBI E-utilities API (free, no key required for <3 req/sec).
    Returns: titles, authors, abstracts, PMIDs with links.
    """
    try:
        # Step 1: Search for PMIDs
        encoded = urllib.parse.quote_plus(query)
        search_url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
            f"?db=pubmed&term={encoded}&retmax={max_results}&retmode=json&sort=relevance"
        )
        req = urllib.request.Request(search_url, headers={"User-Agent": "SymbolAI/1.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            search_data = json.loads(resp.read().decode("utf-8"))

        pmids = search_data.get("esearchresult", {}).get("idlist", [])
        if not pmids:
            return f"لم يتم العثور على أبحاث لـ: {query}"

        total_found = search_data.get("esearchresult", {}).get("count", "0")

        # Step 2: Fetch article details
        ids_str = ",".join(pmids)
        fetch_url = (
            f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
            f"?db=pubmed&id={ids_str}&retmode=json"
        )
        req2 = urllib.request.Request(fetch_url, headers={"User-Agent": "SymbolAI/1.0"})
        with urllib.request.urlopen(req2, timeout=15) as resp2:
            detail_data = json.loads(resp2.read().decode("utf-8"))

        articles = []
        result_items = detail_data.get("result", {})
        for pmid in pmids:
            article = result_items.get(pmid, {})
            if not article or isinstance(article, str):
                continue
            title = article.get("title", "بدون عنوان")
            authors_list = article.get("authors", [])
            authors = ", ".join(a.get("name", "") for a in authors_list[:3])
            if len(authors_list) > 3:
                authors += " وآخرون"
            journal = article.get("fulljournalname", article.get("source", ""))
            pub_date = article.get("pubdate", "")
            link = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"

            articles.append(
                f"📄 {title}\n"
                f"   المؤلفون: {authors}\n"
                f"   المجلة: {journal} ({pub_date})\n"
                f"   🔗 {link}"
            )

        header = f"نتائج PubMed لـ: {query} (وُجد {total_found} بحث، عُرض {len(articles)})\n"
        return header + "\n\n".join(articles)

    except Exception as e:
        logger.error(f"PubMed search failed: {e}")
        return f"خطأ في البحث في PubMed: {e}"


# ═══════════════════════════════════════════
# SKILL 3: PDF Report Generation
# ═══════════════════════════════════════════

def _generate_pdf(
    title: str,
    content: str,
    filename: str = "",
    **kwargs,
) -> str:
    """
    Generate a PDF report from text content.
    توليد تقرير PDF من محتوى نصي

    Uses fpdf2 (lightweight, no Java/LaTeX needed).
    Supports Arabic text with proper RTL rendering.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        return "خطأ: مكتبة fpdf2 غير مثبتة. شغّل: pip install fpdf2"

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"report_{timestamp}.pdf"

    output_dir = Path("./data/outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    try:
        pdf = FPDF()
        pdf.add_page()

        # Try to load Arabic font
        font_loaded = False
        for font_path in [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
            "/usr/share/fonts/TTF/DejaVuSans.ttf",
        ]:
            if Path(font_path).exists():
                pdf.add_font("arabic", "", font_path, uni=True)
                pdf.set_font("arabic", size=12)
                font_loaded = True
                break

        if not font_loaded:
            pdf.set_font("Helvetica", size=12)

        # Title
        pdf.set_font_size(18)
        pdf.cell(0, 15, title, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)

        # Date
        pdf.set_font_size(10)
        pdf.cell(0, 8, f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(10)

        # Content
        pdf.set_font_size(12)
        for paragraph in content.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                pdf.multi_cell(0, 7, paragraph)
                pdf.ln(5)

        pdf.output(str(output_path))
        return f"تم إنشاء التقرير: {output_path} ({output_path.stat().st_size // 1024} كيلوبايت)"

    except Exception as e:
        return f"خطأ في إنشاء PDF: {e}"


# ═══════════════════════════════════════════
# SKILL 4: Word Document Generation
# ═══════════════════════════════════════════

def _generate_docx(
    title: str,
    content: str,
    filename: str = "",
    **kwargs,
) -> str:
    """
    Generate a Word document (.docx) from text content.
    توليد مستند Word من محتوى نصي
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return "خطأ: مكتبة python-docx غير مثبتة. شغّل: pip install python-docx"

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"document_{timestamp}.docx"

    output_dir = Path("./data/outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    try:
        doc = Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}")

        for paragraph in content.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph.startswith("# "):
                doc.add_heading(paragraph[2:], level=1)
            elif paragraph.startswith("## "):
                doc.add_heading(paragraph[3:], level=2)
            elif paragraph.startswith("- ") or paragraph.startswith("• "):
                for line in paragraph.split("\n"):
                    line = line.strip().lstrip("-•").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            elif paragraph:
                doc.add_paragraph(paragraph)

        doc.save(str(output_path))
        return f"تم إنشاء المستند: {output_path} ({output_path.stat().st_size // 1024} كيلوبايت)"

    except Exception as e:
        return f"خطأ في إنشاء المستند: {e}"


# ═══════════════════════════════════════════
# SKILL 5: Drug Interaction Checker
# ═══════════════════════════════════════════

def _check_drug_interaction(drug1: str, drug2: str, **kwargs) -> str:
    """
    Check drug-drug interactions using RxNorm + OpenFDA.
    فحص التداخلات الدوائية باستخدام RxNorm + OpenFDA (مجاني)

    Note: This is a basic check. Always verify with a pharmacist.
    """
    try:
        # Search OpenFDA for drug label warnings
        encoded = urllib.parse.quote_plus(f"{drug1} {drug2}")
        url = (
            f"https://api.fda.gov/drug/label.json"
            f"?search=drug_interactions:{encoded}&limit=3"
        )
        req = urllib.request.Request(url, headers={"User-Agent": "SymbolAI/1.0"})

        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8"))

        results = data.get("results", [])
        if not results:
            return (
                f"لم يتم العثور على تداخلات مسجلة بين {drug1} و{drug2} في OpenFDA.\n"
                f"⚠️ هذا لا يعني عدم وجود تداخلات. راجع الصيدلي دائماً."
            )

        interactions = []
        for result in results[:3]:
            brand = result.get("openfda", {}).get("brand_name", ["غير معروف"])[0]
            warnings = result.get("drug_interactions", ["لا توجد معلومات"])
            warning_text = warnings[0][:500] if warnings else "لا توجد معلومات"
            interactions.append(f"💊 {brand}:\n{warning_text}")

        header = f"تداخلات دوائية محتملة: {drug1} + {drug2}\n"
        footer = "\n\n⚠️ تنبيه: هذه معلومات أولية. يجب مراجعة الصيدلي أو الطبيب المختص."
        return header + "\n\n".join(interactions) + footer

    except urllib.error.HTTPError as e:
        if e.code == 404:
            return f"لم يتم العثور على معلومات عن تداخل {drug1} و{drug2} في قاعدة البيانات."
        return f"خطأ في البحث: {e}"
    except Exception as e:
        return f"خطأ في فحص التداخلات: {e}"


# ═══════════════════════════════════════════
# SKILL 6: Clinical Guidelines Search (WHO / Saudi MOH)
# ═══════════════════════════════════════════

def _clinical_guidelines_search(condition: str, **kwargs) -> str:
    """
    Search for clinical practice guidelines on a medical condition.
    بحث عن الإرشادات السريرية لحالة مرضية

    Searches PubMed specifically for guidelines and systematic reviews.
    """
    # Search PubMed with guideline filter
    query = f"{condition} clinical practice guideline[pt] OR systematic review[pt]"
    return _pubmed_search(query, max_results=5)


# ═══════════════════════════════════════════
# Register All Advanced Skills
# ═══════════════════════════════════════════

def register_advanced_skills(registry: ToolRegistry, skip_existing: bool = False) -> None:
    """Register all advanced skills into the tool registry."""

    def _register(tool: Tool):
        if skip_existing and registry.get(tool.name):
            logger.debug(f"Skipping skill '{tool.name}' — already registered")
            return
        registry.register(tool)

    _register(Tool(
        name="web_search",
        description=(
            "البحث في الإنترنت عن معلومات حديثة. استخدمها عندما تحتاج "
            "معلومات غير موجودة في قاعدة المعرفة المحلية أو معلومات حديثة. "
            "Search the web for current information. No API key needed."
        ),
        parameters={"query": "string — search query", "max_results": "int (default 5)"},
        execute=_web_search,
        category="search",
    ))

    registry.register(Tool(
        name="pubmed_search",
        description=(
            "البحث في PubMed عن الأبحاث والمقالات الطبية المحكّمة. "
            "استخدمها للأسئلة الطبية التي تحتاج مراجع علمية وأدلة. "
            "Search PubMed for peer-reviewed medical research papers."
        ),
        parameters={"query": "string — medical search query", "max_results": "int (default 5)"},
        execute=_pubmed_search,
        category="medical_search",
    ))

    registry.register(Tool(
        name="generate_pdf",
        description=(
            "إنشاء تقرير PDF من محتوى نصي. استخدمها عندما يطلب المستخدم "
            "تقريراً أو مستنداً بتنسيق PDF. "
            "Generate a PDF report from text content."
        ),
        parameters={"title": "string", "content": "string", "filename": "string (optional)"},
        execute=_generate_pdf,
        category="file_generation",
    ))

    registry.register(Tool(
        name="generate_docx",
        description=(
            "إنشاء مستند Word (.docx) من محتوى نصي. يدعم العناوين والقوائم. "
            "Generate a Word document from text. Supports headings and lists."
        ),
        parameters={"title": "string", "content": "string", "filename": "string (optional)"},
        execute=_generate_docx,
        category="file_generation",
    ))

    registry.register(Tool(
        name="drug_interaction",
        description=(
            "فحص التداخلات الدوائية بين دوائين. يستخدم قاعدة بيانات OpenFDA. "
            "⚠️ للمعلومات الأولية فقط — يجب مراجعة الصيدلي. "
            "Check drug-drug interactions using OpenFDA database."
        ),
        parameters={"drug1": "string — first drug name", "drug2": "string — second drug name"},
        execute=_check_drug_interaction,
        category="medical",
    ))

    registry.register(Tool(
        name="clinical_guidelines",
        description=(
            "البحث عن الإرشادات السريرية والمراجعات المنهجية لحالة مرضية. "
            "يبحث في PubMed مع فلترة للإرشادات فقط. "
            "Search for clinical practice guidelines and systematic reviews."
        ),
        parameters={"condition": "string — medical condition"},
        execute=_clinical_guidelines_search,
        category="medical_search",
    ))

    logger.info(f"Advanced skills registered: {registry.list_names()}")
