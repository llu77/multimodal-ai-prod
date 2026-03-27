"""
Extended Agent Tools — Web Search, PubMed, Image Gen, File Gen.
أدوات الوكيل المتقدمة — بحث الويب، PubMed، توليد صور، توليد ملفات

All tools work LOCALLY — no external AI APIs needed.
Web/PubMed use public APIs. Image gen uses local model or SVG.
File gen uses python-docx/reportlab.
"""
import re
import json
import time
import hashlib
from pathlib import Path
from typing import Optional
from datetime import datetime
from loguru import logger


# ═══════════════════════════════════════════
# Web Search
# ═══════════════════════════════════════════

def web_search(query: str, max_results: int = 5, language: str = "ar", **kwargs) -> str:
    """
    Search the web using DuckDuckGo (no API key needed).
    البحث في الويب باستخدام DuckDuckGo (بدون مفتاح API)
    """
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results, region="xa-ar" if language == "ar" else "wt-wt"))

        if not results:
            return f"لم يتم العثور على نتائج لـ: {query}"

        formatted = []
        for i, r in enumerate(results, 1):
            title = r.get("title", "بدون عنوان")
            body = r.get("body", "")[:300]
            url = r.get("href", "")
            formatted.append(f"[{i}] {title}\n{body}\nالرابط: {url}")

        return "\n\n".join(formatted)

    except ImportError:
        return "خطأ: مكتبة duckduckgo_search غير مثبتة. شغّل: pip install duckduckgo-search"
    except Exception as e:
        return f"خطأ في البحث: {str(e)}"


# ═══════════════════════════════════════════
# PubMed Search
# ═══════════════════════════════════════════

def pubmed_search(query: str, max_results: int = 5, **kwargs) -> str:
    """
    Search PubMed for medical research papers.
    البحث في PubMed عن أبحاث طبية

    Uses NCBI E-utilities API (free, no key needed for <3 requests/sec).
    """
    import urllib.request
    import urllib.parse
    import xml.etree.ElementTree as ET

    base_search = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    base_fetch = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"

    try:
        # Step 1: Search for article IDs
        search_params = urllib.parse.urlencode({
            "db": "pubmed",
            "term": query,
            "retmax": max_results,
            "sort": "relevance",
            "retmode": "json",
        })
        search_url = f"{base_search}?{search_params}"

        with urllib.request.urlopen(search_url, timeout=15) as resp:
            search_data = json.loads(resp.read().decode("utf-8"))

        id_list = search_data.get("esearchresult", {}).get("idlist", [])
        if not id_list:
            return f"لم يتم العثور على أبحاث لـ: {query}"

        # Step 2: Fetch article details
        fetch_params = urllib.parse.urlencode({
            "db": "pubmed",
            "id": ",".join(id_list),
            "retmode": "xml",
            "rettype": "abstract",
        })
        fetch_url = f"{base_fetch}?{fetch_params}"

        with urllib.request.urlopen(fetch_url, timeout=15) as resp:
            xml_data = resp.read().decode("utf-8")

        root = ET.fromstring(xml_data)
        articles = []

        for article in root.findall(".//PubmedArticle"):
            # Title
            title_el = article.find(".//ArticleTitle")
            title = title_el.text if title_el is not None and title_el.text else "بدون عنوان"

            # Authors
            authors = []
            for author in article.findall(".//Author")[:3]:
                last = author.find("LastName")
                first = author.find("ForeName")
                if last is not None and last.text:
                    name = last.text
                    if first is not None and first.text:
                        name += f" {first.text[0]}"
                    authors.append(name)
            author_str = "، ".join(authors)
            if len(article.findall(".//Author")) > 3:
                author_str += " وآخرون"

            # Year
            year_el = article.find(".//PubDate/Year")
            year = year_el.text if year_el is not None else "—"

            # Journal
            journal_el = article.find(".//Journal/Title")
            journal = journal_el.text if journal_el is not None and journal_el.text else ""

            # Abstract
            abstract_parts = []
            for abs_text in article.findall(".//AbstractText"):
                label = abs_text.get("Label", "")
                text = abs_text.text or ""
                if label:
                    abstract_parts.append(f"{label}: {text[:200]}")
                else:
                    abstract_parts.append(text[:300])
            abstract = " ".join(abstract_parts) if abstract_parts else "لا يوجد ملخص"

            # PMID
            pmid_el = article.find(".//PMID")
            pmid = pmid_el.text if pmid_el is not None else ""
            link = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else ""

            articles.append(
                f"📄 {title}\n"
                f"   المؤلفون: {author_str} ({year})\n"
                f"   المجلة: {journal}\n"
                f"   الملخص: {abstract}\n"
                f"   الرابط: {link}"
            )

        return "\n\n".join(articles)

    except Exception as e:
        return f"خطأ في البحث في PubMed: {str(e)}"


# ═══════════════════════════════════════════
# Image Generation (SVG — no GPU needed)
# ═══════════════════════════════════════════

def generate_rehab_image(
    exercise_type: str,
    body_part: str = "",
    title: str = "",
    output_dir: str = "./data/generated",
    **kwargs,
) -> str:
    """
    Generate rehabilitation exercise illustrations as SVG.
    توليد رسومات تمارين إعادة التأهيل كـ SVG

    No GPU needed — uses programmatic SVG generation.
    Produces clean medical-style exercise diagrams.
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # Exercise SVG templates
    templates = {
        "shoulder_pendulum": _svg_shoulder_pendulum,
        "shoulder_flexion": _svg_shoulder_flexion,
        "knee_extension": _svg_knee_extension,
        "back_extension": _svg_back_extension,
        "general_stretch": _svg_general_stretch,
    }

    # Match exercise type
    exercise_key = exercise_type.lower().replace(" ", "_")
    # Fuzzy match
    matched = None
    for key in templates:
        if exercise_key in key or key in exercise_key:
            matched = key
            break
    if not matched:
        # Default
        for key in templates:
            if body_part and body_part in key:
                matched = key
                break
    if not matched:
        matched = "general_stretch"

    svg_content = templates[matched](title or exercise_type)

    # Save
    filename = f"exercise_{matched}_{int(time.time())}.svg"
    filepath = output_path / filename
    filepath.write_text(svg_content, encoding="utf-8")

    return f"تم توليد صورة التمرين: {filepath}\nالنوع: {matched}\nالتنسيق: SVG (قابل للطباعة بأي حجم)"


def _svg_shoulder_pendulum(title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="400" height="400">
  <rect width="400" height="400" fill="#f8f9fa" rx="12"/>
  <text x="200" y="30" text-anchor="middle" font-family="Arial" font-size="16" font-weight="bold" fill="#2c3e50">{title}</text>
  <!-- Table -->
  <rect x="180" y="80" width="200" height="20" fill="#8B4513" rx="3"/>
  <rect x="350" y="100" width="20" height="200" fill="#8B4513" rx="3"/>
  <!-- Body leaning on table -->
  <circle cx="280" cy="120" r="25" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <line x1="280" y1="145" x2="280" y2="250" stroke="#333" stroke-width="3"/>
  <line x1="280" y1="250" x2="260" y2="350" stroke="#333" stroke-width="3"/>
  <line x1="280" y1="250" x2="300" y2="350" stroke="#333" stroke-width="3"/>
  <!-- Support arm on table -->
  <line x1="280" y1="170" x2="340" y2="95" stroke="#333" stroke-width="3"/>
  <!-- Pendulum arm hanging -->
  <line x1="280" y1="170" x2="220" y2="280" stroke="#333" stroke-width="3"/>
  <!-- Swing arc -->
  <path d="M 190 280 Q 220 310 250 280" fill="none" stroke="#e74c3c" stroke-width="2" stroke-dasharray="5,5"/>
  <polygon points="188,275 185,285 195,280" fill="#e74c3c"/>
  <polygon points="252,275 255,285 245,280" fill="#e74c3c"/>
  <!-- Label -->
  <text x="200" y="380" text-anchor="middle" font-family="Arial" font-size="12" fill="#555">تمرين البندول — حرّك الذراع كالبندول برفق</text>
</svg>'''


def _svg_shoulder_flexion(title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="400" height="400">
  <rect width="400" height="400" fill="#f8f9fa" rx="12"/>
  <text x="200" y="30" text-anchor="middle" font-family="Arial" font-size="16" font-weight="bold" fill="#2c3e50">{title}</text>
  <!-- Standing figure -->
  <circle cx="200" cy="80" r="25" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <line x1="200" y1="105" x2="200" y2="230" stroke="#333" stroke-width="3"/>
  <line x1="200" y1="230" x2="180" y2="350" stroke="#333" stroke-width="3"/>
  <line x1="200" y1="230" x2="220" y2="350" stroke="#333" stroke-width="3"/>
  <!-- Right arm down -->
  <line x1="200" y1="140" x2="160" y2="230" stroke="#333" stroke-width="3"/>
  <!-- Left arm raising (with arc) -->
  <line x1="200" y1="140" x2="230" y2="70" stroke="#2ecc71" stroke-width="3"/>
  <path d="M 240 230 Q 260 150 230 70" fill="none" stroke="#e74c3c" stroke-width="2" stroke-dasharray="5,5"/>
  <polygon points="228,68 235,65 232,75" fill="#e74c3c"/>
  <text x="200" y="380" text-anchor="middle" font-family="Arial" font-size="12" fill="#555">رفع الذراع أماماً — ارفع ببطء حتى مستوى الرأس</text>
</svg>'''


def _svg_knee_extension(title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="400" height="400">
  <rect width="400" height="400" fill="#f8f9fa" rx="12"/>
  <text x="200" y="30" text-anchor="middle" font-family="Arial" font-size="16" font-weight="bold" fill="#2c3e50">{title}</text>
  <!-- Chair -->
  <rect x="80" y="150" width="120" height="15" fill="#8B4513" rx="3"/>
  <rect x="80" y="165" width="15" height="180" fill="#8B4513"/>
  <rect x="185" y="165" width="15" height="180" fill="#8B4513"/>
  <rect x="80" y="80" width="15" height="85" fill="#8B4513"/>
  <!-- Seated figure -->
  <circle cx="160" cy="90" r="22" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <line x1="160" y1="112" x2="160" y2="165" stroke="#333" stroke-width="3"/>
  <!-- Thigh on chair -->
  <line x1="160" y1="165" x2="210" y2="200" stroke="#333" stroke-width="3"/>
  <!-- Lower leg extending -->
  <line x1="210" y1="200" x2="300" y2="200" stroke="#2ecc71" stroke-width="3"/>
  <!-- Extension arc -->
  <path d="M 210 200 Q 250 280 210 260" fill="none" stroke="#aaa" stroke-width="2" stroke-dasharray="4,4"/>
  <line x1="210" y1="200" x2="210" y2="260" stroke="#aaa" stroke-width="2"/>
  <path d="M 250 260 Q 280 230 300 200" fill="none" stroke="#e74c3c" stroke-width="2" stroke-dasharray="5,5"/>
  <polygon points="302,198 296,205 305,204" fill="#e74c3c"/>
  <text x="200" y="380" text-anchor="middle" font-family="Arial" font-size="12" fill="#555">بسط الركبة — ابسط ببطء ثم أنزل ببطء</text>
</svg>'''


def _svg_back_extension(title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="400" height="400">
  <rect width="400" height="400" fill="#f8f9fa" rx="12"/>
  <text x="200" y="30" text-anchor="middle" font-family="Arial" font-size="16" font-weight="bold" fill="#2c3e50">{title}</text>
  <!-- Floor -->
  <line x1="50" y1="320" x2="350" y2="320" stroke="#aaa" stroke-width="1"/>
  <!-- Prone figure - starting -->
  <ellipse cx="200" cy="300" rx="80" ry="15" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <circle cx="120" cy="295" r="15" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <!-- Arms supporting -->
  <line x1="140" y1="300" x2="150" y2="315" stroke="#333" stroke-width="2"/>
  <line x1="160" y1="300" x2="170" y2="315" stroke="#333" stroke-width="2"/>
  <!-- Extension arc showing upper body lift -->
  <path d="M 120 295 Q 100 250 120 220" fill="none" stroke="#e74c3c" stroke-width="2" stroke-dasharray="5,5"/>
  <polygon points="118,218 125,222 115,225" fill="#e74c3c"/>
  <!-- Extended position (ghost) -->
  <ellipse cx="200" cy="280" rx="80" ry="12" fill="none" stroke="#2ecc71" stroke-width="1.5" stroke-dasharray="4,4"/>
  <circle cx="120" cy="240" r="15" fill="none" stroke="#2ecc71" stroke-width="1.5" stroke-dasharray="4,4"/>
  <text x="200" y="380" text-anchor="middle" font-family="Arial" font-size="12" fill="#555">تمديد الظهر — ارفع الجزء العلوي ببطء مع الاستناد على اليدين</text>
</svg>'''


def _svg_general_stretch(title: str) -> str:
    return f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="400" height="400">
  <rect width="400" height="400" fill="#f8f9fa" rx="12"/>
  <text x="200" y="30" text-anchor="middle" font-family="Arial" font-size="16" font-weight="bold" fill="#2c3e50">{title}</text>
  <!-- Standing figure -->
  <circle cx="200" cy="90" r="25" fill="#FDBCB4" stroke="#333" stroke-width="2"/>
  <line x1="200" y1="115" x2="200" y2="240" stroke="#333" stroke-width="3"/>
  <line x1="200" y1="240" x2="175" y2="350" stroke="#333" stroke-width="3"/>
  <line x1="200" y1="240" x2="225" y2="350" stroke="#333" stroke-width="3"/>
  <!-- Arms stretching up -->
  <line x1="200" y1="145" x2="140" y2="80" stroke="#2ecc71" stroke-width="3"/>
  <line x1="200" y1="145" x2="260" y2="80" stroke="#2ecc71" stroke-width="3"/>
  <!-- Stretch arrows -->
  <polygon points="137,78 145,72 143,83" fill="#e74c3c"/>
  <polygon points="263,78 255,72 257,83" fill="#e74c3c"/>
  <text x="200" y="380" text-anchor="middle" font-family="Arial" font-size="12" fill="#555">تمرين إطالة عامة — ارفع الذراعين واثبت 15 ثانية</text>
</svg>'''


# ═══════════════════════════════════════════
# File / Report Generation
# ═══════════════════════════════════════════

def generate_report(
    title: str,
    content: str,
    output_format: str = "txt",
    patient_name: str = "",
    output_dir: str = "./data/generated",
    **kwargs,
) -> str:
    """
    Generate a medical/rehab report file.
    توليد تقرير طبي/تأهيلي كملف

    Formats: txt (always works), docx (needs python-docx), pdf (needs reportlab)
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip().replace(' ', '_')

    if output_format == "docx":
        return _generate_docx(output_path, safe_title, timestamp, title, content, patient_name)
    elif output_format == "pdf":
        return _generate_pdf(output_path, safe_title, timestamp, title, content, patient_name)
    else:
        return _generate_txt(output_path, safe_title, timestamp, title, content, patient_name)


def _generate_txt(output_path, safe_title, timestamp, title, content, patient_name) -> str:
    filename = f"report_{safe_title}_{timestamp}.txt"
    filepath = output_path / filename

    lines = [
        "=" * 60,
        f"  {title}",
        "=" * 60,
        f"  التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    ]
    if patient_name:
        lines.append(f"  المريض: {patient_name}")
    lines.extend([
        "=" * 60,
        "",
        content,
        "",
        "=" * 60,
        "  تم التوليد بواسطة Symbol Rehab AI",
        "=" * 60,
    ])

    filepath.write_text("\n".join(lines), encoding="utf-8")
    return f"تم توليد التقرير: {filepath}\nالتنسيق: نص عادي (TXT)"


def _generate_docx(output_path, safe_title, timestamp, title, content, patient_name) -> str:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta.add_run(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(10)
        if patient_name:
            meta.add_run(f"\nالمريض: {patient_name}").font.size = Pt(10)

        doc.add_paragraph("")  # Spacer

        # Content — split by paragraphs
        for paragraph in content.split("\n"):
            p = doc.add_paragraph(paragraph.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(6)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph("تم التوليد بواسطة Symbol Rehab AI")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

        filename = f"report_{safe_title}_{timestamp}.docx"
        filepath = output_path / filename
        doc.save(str(filepath))
        return f"تم توليد التقرير: {filepath}\nالتنسيق: مستند Word (DOCX)"

    except ImportError:
        logger.warning("python-docx not installed, falling back to TXT")
        return _generate_txt(output_path, safe_title, timestamp, title, content, patient_name)


def _generate_pdf(output_path, safe_title, timestamp, title, content, patient_name) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm

        filename = f"report_{safe_title}_{timestamp}.pdf"
        filepath = output_path / filename

        c = canvas.Canvas(str(filepath), pagesize=A4)
        width, height = A4

        # Title
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(width / 2, height - 3 * cm, title)

        # Metadata
        c.setFont("Helvetica", 10)
        y = height - 4.5 * cm
        c.drawString(2 * cm, y, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if patient_name:
            y -= 0.5 * cm
            c.drawString(2 * cm, y, f"Patient: {patient_name}")

        # Content
        c.setFont("Helvetica", 11)
        y -= 1.5 * cm
        for line in content.split("\n"):
            if y < 3 * cm:
                c.showPage()
                y = height - 3 * cm
                c.setFont("Helvetica", 11)
            c.drawString(2 * cm, y, line.strip()[:90])
            y -= 0.5 * cm

        # Footer
        c.setFont("Helvetica", 8)
        c.drawCentredString(width / 2, 1.5 * cm, "Generated by Symbol Rehab AI")

        c.save()
        return f"تم توليد التقرير: {filepath}\nالتنسيق: PDF"

    except ImportError:
        logger.warning("reportlab not installed, falling back to TXT")
        return _generate_txt(output_path, safe_title, timestamp, title, content, patient_name)
